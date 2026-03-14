from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import FileResponse

import os

app = FastAPI()

@app.get("/robots.txt")
def robots():
    return FileResponse("static/robots.txt", media_type="text/plain")


@app.get("/sitemap.xml")
def sitemap():
    return FileResponse("sitemap.xml", media_type="application/xml")

# Static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Templates
templates = Jinja2Templates(directory="templates")


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


@app.post("/upload")
async def upload(
    file: UploadFile = File(...),
    margenes: bool = Form(False),
    fuente: bool = Form(False),
    interlineado: bool = Form(False),
    sangria: bool = Form(False),
    numeracion: bool = Form(False),
    referencias: bool = Form(False)
):

    contents = await file.read()

    doc = Document(BytesIO(contents))

    # Aplicar márgenes
    if margenes:
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Procesar párrafos
    in_references = False

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()
        style_name = paragraph.style.name.lower()

        # detectar sección de referencias
        if text.lower() == "referencias":
            in_references = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # evitar modificar párrafos con imágenes
        if paragraph._element.xpath('.//pic:pic'):
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if fuente:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        if interlineado:
            paragraph.paragraph_format.line_spacing = 2

        # detectar si es título o lista
        es_titulo = "heading" in style_name or "título" in style_name
        es_lista = "list" in style_name

        # aplicar sangría solo a texto normal
        if sangria and not es_titulo and not es_lista and not in_references:
            paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # referencias → sangría francesa
        if referencias and in_references:
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)

    # Numeración de páginas (FUERA del loop)
    if numeracion:

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for section in doc.sections:

            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.clear()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = paragraph.add_run()

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

    # Guardar en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    original_name = file.filename.replace(".docx", "")

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={original_name}_APA.docx"
        }
    )

@app.get("/privacy", response_class=HTMLResponse)
async def privacy(request: Request):
    return templates.TemplateResponse("privacy.html", {"request": request})


@app.get("/terms", response_class=HTMLResponse)
async def terms(request: Request):
    return templates.TemplateResponse("terms.html", {"request": request})


@app.get("/contact", response_class=HTMLResponse)
async def contact(request: Request):
    return templates.TemplateResponse("contact.html", {"request": request})